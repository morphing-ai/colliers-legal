"""
DOCX Export Service for Compliance Analysis Results
Exports compliance analysis with violations and recommendations to Word format using native python-docx comment support
"""
import io
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import os

logger = logging.getLogger(__name__)


class DocxExportService:
    """Service for exporting compliance analysis results to DOCX format using native python-docx comments."""
    
    def __init__(self):
        self.severity_colors = {
            'critical': RGBColor(220, 53, 69),   # Red
            'high': RGBColor(253, 126, 20),      # Orange
            'medium': RGBColor(255, 193, 7),     # Yellow
            'low': RGBColor(13, 110, 253),       # Blue
            'success': RGBColor(25, 135, 84)     # Green
        }
    
    def _format_violation_comment(self, issue: Dict[str, Any]) -> str:
        """Format a violation issue into a comprehensive comment text."""
        lines = []
        
        # Severity header
        severity = issue.get('severity', 'medium').upper()
        lines.append(f"🚨 {severity} COMPLIANCE ISSUE")
        lines.append("=" * 40)
        
        # Rule information
        if issue.get('rule_number'):
            lines.append(f"📋 Rule: {issue['rule_number']}")
        if issue.get('rule_title'):
            lines.append(f"📖 Title: {issue['rule_title']}")
        if issue.get('issue_type'):
            lines.append(f"🏷️ Type: {issue['issue_type']}")
        
        lines.append("")
        
        # Issue description
        if issue.get('description'):
            lines.append("📝 ISSUE DESCRIPTION:")
            lines.append(issue['description'])
            lines.append("")
        
        # Current vs Required text
        if issue.get('current_text'):
            lines.append("❌ CURRENT TEXT:")
            lines.append(issue['current_text'])
            lines.append("")
        
        if issue.get('required_text'):
            lines.append("✅ REQUIRED TEXT:")
            lines.append(issue['required_text'])
            lines.append("")
        
        # Recommended action
        if issue.get('suggested_fix'):
            lines.append("🔧 RECOMMENDED ACTION:")
            lines.append(issue['suggested_fix'])
        
        return "\n".join(lines)
    
    def _add_native_comment(self, doc: Document, paragraph, runs: List, comment_text: str, author: str = "Compliance Analyzer") -> bool:
        """Add a native comment to specified runs using python-docx 1.2.0+ comment support."""
        try:
            # Use native comment support
            comment = doc.add_comment(
                runs=runs,
                text=comment_text,
                author=author,
                initials="CA"
            )
            return True
        except Exception as e:
            logger.warning(f"Failed to add native comment: {e}. Falling back to inline annotation.")
            return False
    
    def _add_violation_comment_to_paragraph(self, doc: Document, paragraph, issue: Dict[str, Any]) -> bool:
        """Add a comment for a violation to an entire paragraph."""
        if not paragraph.runs:
            # If no runs, create one with the paragraph text
            paragraph.add_run(paragraph.text or "[No text]")
        
        # Format the comprehensive comment
        comment_text = self._format_violation_comment(issue)
        
        # Try to add native comment
        success = self._add_native_comment(
            doc=doc,
            paragraph=paragraph,
            runs=paragraph.runs,
            comment_text=comment_text,
            author="Compliance Analyzer"
        )
        
        if not success:
            # Fallback: Add inline annotation
            severity = issue.get('severity', 'medium')
            inline_text = f" [{severity.upper()}: {issue.get('rule_number', 'N/A')} - {issue.get('description', 'Compliance issue')}]"
            
            comment_run = paragraph.add_run(inline_text)
            comment_run.font.color.rgb = self.severity_colors.get(severity, RGBColor(128, 128, 128))
            comment_run.font.size = Pt(9)
            comment_run.font.italic = True
        
        return success
    
    async def export_analysis(
        self, 
        analysis_results: Dict[str, Any],
        rule_set_name: str = "FINRA Rules"
    ) -> bytes:
        """
        Export compliance analysis results to DOCX format with native comments.
        
        Args:
            analysis_results: Analysis results from the API
            rule_set_name: Name of the rule set used for analysis
            
        Returns:
            DOCX file as bytes
        """        
        # Create a new document
        doc = Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Add title page
        self._add_title_page(doc, analysis_results, rule_set_name)
        
        # Add executive summary
        self._add_executive_summary(doc, analysis_results)
        
        # Add detailed findings with native comments
        self._add_detailed_findings_with_comments(doc, analysis_results)
        
        # Add recommendations summary
        self._add_recommendations_summary(doc, analysis_results)
        
        # Save to temporary file and return as bytes
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            temp_path = tmp_file.name
        
        try:
            # Read the file as bytes
            with open(temp_path, 'rb') as f:
                doc_bytes = f.read()
            
            return doc_bytes
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
    
    def _setup_styles(self, doc):
        """Set up custom styles for the document."""
        styles = doc.styles
        
        # Create heading styles if they don't exist
        if 'Custom Heading 1' not in styles:
            heading_1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading_1.font.size = Pt(18)
            heading_1.font.bold = True
            heading_1.font.color.rgb = RGBColor(0, 0, 0)
            heading_1.paragraph_format.space_after = Pt(12)
    
    def _add_title_page(self, doc, results, rule_set_name):
        """Add title page to the document."""
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("Compliance Analysis Report")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        doc.add_paragraph()
        
        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(f"Analysis against {rule_set_name}")
        subtitle_run.font.size = Pt(16)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_run = meta_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
        date_run.font.size = Pt(12)
        
        # Session ID
        if results.get('session_id'):
            session_run = meta_para.add_run(f"Session ID: {results['session_id']}\n")
            session_run.font.size = Pt(10)
            session_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Page break
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, results):
        """Add executive summary section."""
        doc.add_heading('Executive Summary', level=1)
        
        paragraphs = results.get('paragraphs', [])
        total_issues = sum(len(p.get('issues', [])) for p in paragraphs)
        
        # Count by severity
        severity_counts = {
            'critical': 0,
            'high': 0,
            'medium': 0,
            'low': 0
        }
        
        for para in paragraphs:
            for issue in para.get('issues', []):
                severity = issue.get('severity', 'medium')
                if severity in severity_counts:
                    severity_counts[severity] += 1
        
        # Summary paragraph
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(
            f"The compliance analysis examined {len(paragraphs)} paragraphs and identified {total_issues} compliance issues. "
        )
        summary_run.font.size = Pt(12)
        
        # Add statistics table
        doc.add_heading('Issue Summary', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Severity'
        header_cells[1].text = 'Count'
        
        # Data rows
        for severity, count in severity_counts.items():
            row_cells = table.add_row().cells
            row_cells[0].text = severity.capitalize()
            row_cells[1].text = str(count)
            
            # Color code the severity cell
            if count > 0:
                run = row_cells[0].paragraphs[0].runs[0]
                run.font.bold = True
                if severity in self.severity_colors:
                    run.font.color.rgb = self.severity_colors[severity]
        
        doc.add_paragraph()  # Spacing
    
    def _add_detailed_findings_with_comments(self, doc, results):
        """Add detailed findings section with native comments for violations."""
        doc.add_heading('Detailed Findings', level=1)
        
        paragraphs = results.get('paragraphs', [])
        comment_count = 0
        
        for idx, para in enumerate(paragraphs):
            issues = para.get('issues', [])
            
            # Add paragraph header
            doc.add_heading(f'Paragraph {para.get("index", idx) + 1}', level=2)
            
            # Add the original text with comments for violations
            doc.add_heading('Document Text:', level=3)
            text_para = doc.add_paragraph(para.get('content', ''))
            text_para.style = 'Quote'
            
            # Add comments for each issue found in this paragraph
            if issues:
                for issue in issues:
                    success = self._add_violation_comment_to_paragraph(doc, text_para, issue)
                    if success:
                        comment_count += 1
                        # Apply severity-based highlighting to the paragraph
                        severity = issue.get('severity', 'medium')
                        for run in text_para.runs:
                            if severity == 'critical':
                                run.font.highlight_color = 2  # Red
                            elif severity == 'high':
                                run.font.highlight_color = 6  # Orange/Red
                            elif severity == 'medium':
                                run.font.highlight_color = 7  # Yellow
                            elif severity == 'low':
                                run.font.highlight_color = 11  # Light Blue
                            break  # Only highlight once per paragraph
            
            # Add a summary of issues for this paragraph
            if issues:
                doc.add_paragraph()  # Spacing
                summary_para = doc.add_paragraph()
                summary_run = summary_para.add_run(f"Found {len(issues)} compliance issue(s) in this paragraph. ")
                summary_run.font.size = Pt(10)
                summary_run.font.italic = True
                summary_run.font.color.rgb = RGBColor(128, 128, 128)
                
                summary_run2 = summary_para.add_run("See comments for detailed analysis and recommendations.")
                summary_run2.font.size = Pt(10)
                summary_run2.font.italic = True
                summary_run2.font.color.rgb = RGBColor(0, 102, 204)
            
            # Add separator between paragraphs
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
        
        # Add comment summary
        if comment_count > 0:
            doc.add_paragraph()
            comment_summary = doc.add_paragraph()
            comment_summary_run = comment_summary.add_run(f"💬 Total Comments Added: {comment_count}")
            comment_summary_run.font.bold = True
            comment_summary_run.font.color.rgb = RGBColor(0, 102, 204)
            
            instruction = doc.add_paragraph()
            instruction_run = instruction.add_run("To view comments in Microsoft Word: Go to Review → Comments, or look for comment indicators in the document margins.")
            instruction_run.font.size = Pt(10)
            instruction_run.font.italic = True
            instruction_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_detailed_findings(self, doc, results):
        """Legacy method - kept for backward compatibility. Use _add_detailed_findings_with_comments instead."""
        return self._add_detailed_findings_with_comments(doc, results)
    
    def _add_recommendations_summary(self, doc, results):
        """Add a summary of all recommendations at the end."""
        doc.add_page_break()
        doc.add_heading('Summary of Recommendations', level=1)
        
        doc.add_paragraph(
            "The following actions are recommended to achieve compliance:"
        )
        
        paragraphs = results.get('paragraphs', [])
        recommendation_count = 0
        
        for para in paragraphs:
            for issue in para.get('issues', []):
                if issue.get('suggested_fix'):
                    recommendation_count += 1
                    
                    # Add recommendation with paragraph reference
                    rec_para = doc.add_paragraph(style='List Bullet')
                    
                    # Paragraph reference
                    ref_run = rec_para.add_run(f"[Paragraph {para.get('index', 0) + 1}] ")
                    ref_run.font.bold = True
                    ref_run.font.size = Pt(10)
                    
                    # Rule reference
                    if issue.get('rule_number'):
                        rule_run = rec_para.add_run(f"Rule {issue['rule_number']}: ")
                        rule_run.font.italic = True
                        rule_run.font.size = Pt(10)
                    
                    # The recommendation
                    rec_para.add_run(issue['suggested_fix'])
        
        if recommendation_count == 0:
            doc.add_paragraph("No specific recommendations were identified.")
        else:
            doc.add_paragraph()
            conclusion = doc.add_paragraph()
            conclusion.add_run(f"Total Recommendations: {recommendation_count}").bold = True
        
        # Add footer
        doc.add_paragraph()
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Generated by Compliance Analyzer")
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
